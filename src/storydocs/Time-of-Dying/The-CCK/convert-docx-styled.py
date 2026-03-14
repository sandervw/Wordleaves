"""
Convert a styled .docx file to Markdown, preserving distinct formatting
for each paragraph style.

Usage:
    python convert-docx-styled.py <input.docx> [output.md] [--media-dir DIR]

Requires: python-docx (pip install python-docx)

Style-to-Markdown mapping (configurable via STYLE_MAP):
    Heading 2          -> ## Heading
    DESC HEAD          -> ### Sub-heading
    DESCRIPTION        -> Regular text
    ACTION             -> `Backtick text`
    DIALOGUE           -> > Blockquote
    REFLECTION         -> > *Italic blockquote*
    Bergamot dividers  -> --- horizontal rule
"""

import sys
import os
import re
import shutil
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def build_style_name_map(doc):
    """Build a dict mapping style_id -> style_name from the document's styles."""
    id_to_name = {}
    for style in doc.styles:
        if style.style_id and style.name:
            id_to_name[style.style_id] = style.name
    return id_to_name

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

STYLE_MAP = {
    "Heading 2":    "heading",
    "DESC HEAD":    "subheading",
    "DESCRIPTION":  "description",
    "ACTION":       "action",
    "DIALOGUE":     "dialogue",
    "REFLECTION":   "reflection",
    "ListStyle":    "list",
}

# The ornamental divider text used with Bergamot Ornaments font
DIVIDER_PATTERN = re.compile(r"^[`a]+$")

# Namespaces for raw XML image extraction
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

def extract_images(docx_path, media_dir):
    """Extract images from docx and return a map of rId -> relative file path."""
    media_dir = Path(media_dir)
    media_dir.mkdir(parents=True, exist_ok=True)

    rId_to_file = {}

    with zipfile.ZipFile(docx_path) as z:
        # Parse document.xml.rels to find image relationship IDs
        rels_xml = z.read("word/_rels/document.xml.rels")
        rels_root = ET.fromstring(rels_xml)
        for rel in rels_root.findall("{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"):
            rel_type = rel.get("Type", "")
            if "image" in rel_type:
                rId = rel.get("Id")
                target = rel.get("Target")  # e.g. "media/image1.png"
                filename = os.path.basename(target)
                # Extract the image file
                src_path = f"word/{target}"
                if src_path in z.namelist():
                    dest = media_dir / filename
                    with z.open(src_path) as src, open(dest, "wb") as dst:
                        shutil.copyfileobj(src, dst)
                    rId_to_file[rId] = str(media_dir / filename).replace("\\", "/")

    return rId_to_file


def find_image_rids(element):
    """Find all image relationship IDs (rId) in a paragraph element."""
    rids = []
    blips = element.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    )
    for blip in blips:
        embed = blip.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        )
        if embed:
            rids.append(embed)
    return rids


# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------

def get_paragraph_text(paragraph):
    """Extract text from a paragraph, preserving inline italic/bold."""
    parts = []
    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue
        if run.italic:
            text = f"*{text}*"
        if run.bold:
            text = f"**{text}**"
        parts.append(text)
    return "".join(parts)


# ---------------------------------------------------------------------------
# Table conversion
# ---------------------------------------------------------------------------

def is_reflection_table(table_element):
    """Check if a table is a single-cell reflection/narrative block."""
    rows = table_element.findall(qn("w:tr"))
    if len(rows) != 1:
        return False
    cells = rows[0].findall(qn("w:tc"))
    # Account for gridSpan -- count logical cells
    logical_cells = 0
    for cell in cells:
        tcpr = cell.find(qn("w:tcPr"))
        if tcpr is not None:
            gs = tcpr.find(qn("w:gridSpan"))
            if gs is not None:
                logical_cells += int(gs.get(qn("w:val"), "1"))
            else:
                logical_cells += 1
        else:
            logical_cells += 1
    return logical_cells == 1


def convert_reflection_table(table_element, style_id_to_name):
    """Convert a single-cell narrative table to bold blockquote lines."""
    lines = []
    rows = table_element.findall(qn("w:tr"))
    cell = rows[0].findall(qn("w:tc"))[0]
    from docx.table import _Cell
    cell_obj = _Cell(cell, None)
    for para in cell_obj.paragraphs:
        text = get_paragraph_text(para)
        if not text.strip():
            lines.append(">")
        else:
            # Check style of para inside cell
            style_id = ""
            ppr = para._element.find(qn("w:pPr"))
            if ppr is not None:
                ps = ppr.find(qn("w:pStyle"))
                if ps is not None:
                    style_id = ps.get(qn("w:val"), "")

            style_name = style_id_to_name.get(style_id, style_id)
            role = STYLE_MAP.get(style_name, "reflection")
            if role == "dialogue":
                lines.append(f"> *{text}*")
            else:
                lines.append(f"> *{text}*")
    return lines


def convert_data_table(table_element, style_id_to_name):
    """Convert a multi-cell table to a markdown pipe table."""
    from docx.table import _Cell

    rows_el = table_element.findall(qn("w:tr"))

    # Build grid: list of rows, each row is list of cell texts
    grid = []
    col_count = 0

    for row_el in rows_el:
        cells_el = row_el.findall(qn("w:tc"))
        row_data = []
        for cell_el in cells_el:
            tcpr = cell_el.find(qn("w:tcPr"))
            span = 1
            if tcpr is not None:
                gs = tcpr.find(qn("w:gridSpan"))
                if gs is not None:
                    span = int(gs.get(qn("w:val"), "1"))

            cell_obj = _Cell(cell_el, None)
            # Collect paragraphs, separate DESC HEAD from DESCRIPTION
            cell_parts = []
            for para in cell_obj.paragraphs:
                text = get_paragraph_text(para).strip()
                if not text:
                    continue
                style_id = ""
                ppr = para._element.find(qn("w:pPr"))
                if ppr is not None:
                    ps = ppr.find(qn("w:pStyle"))
                    if ps is not None:
                        style_id = ps.get(qn("w:val"), "")
                style_name = style_id_to_name.get(style_id, style_id)
                role = STYLE_MAP.get(style_name, "description")
                if role == "subheading":
                    cell_parts.append(f"**{text}**")
                else:
                    cell_parts.append(text)

            cell_text = " ".join(cell_parts)

            # Add cell for each spanned column
            for s in range(span):
                row_data.append(cell_text if s == 0 else "")

            col_count = max(col_count, len(row_data))

        grid.append(row_data)

    # Pad rows to col_count
    for row in grid:
        while len(row) < col_count:
            row.append("")

    # Build markdown table
    lines = []
    for ri, row in enumerate(grid):
        line = "| " + " | ".join(cell if cell else " " for cell in row) + " |"
        lines.append(line)
        if ri == 0:
            sep = "| " + " | ".join("---" for _ in row) + " |"
            lines.append(sep)

    return lines


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------

def convert_docx_to_md(docx_path, output_path=None, media_dir=None):
    docx_path = Path(docx_path)

    if output_path is None:
        output_path = docx_path.with_suffix(".md")
    else:
        output_path = Path(output_path)

    if media_dir is None:
        media_dir = output_path.parent / (output_path.stem + "-media")

    # Extract images
    rId_map = extract_images(docx_path, media_dir)

    doc = Document(str(docx_path))
    style_id_to_name = build_style_name_map(doc)
    md_lines = []

    def add_blank():
        """Add a blank line if the last line isn't already blank."""
        if md_lines and md_lines[-1] != "":
            md_lines.append("")

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            # Get style
            style_id = ""
            ppr = element.find(qn("w:pPr"))
            if ppr is not None:
                ps = ppr.find(qn("w:pStyle"))
                if ps is not None:
                    style_id = ps.get(qn("w:val"), "")

            # Get paragraph object for text extraction
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break

            if para is None:
                # Paragraph might be unreachable via doc.paragraphs (e.g. in headers)
                continue

            text = get_paragraph_text(para)

            # Check for images in this paragraph
            image_rids = find_image_rids(element)
            if image_rids:
                add_blank()
                for rid in image_rids:
                    img_path = rId_map.get(rid, f"[image:{rid}]")
                    md_lines.append(f"![{text.strip()}]({img_path})")
                add_blank()
                # If paragraph only contains the image, skip text output
                if not text.strip():
                    continue

            # Check for ornamental divider
            if DIVIDER_PATTERN.match(text.strip()):
                add_blank()
                md_lines.append("---")
                add_blank()
                continue

            # Skip empty paragraphs
            if not text.strip():
                continue

            style_name = style_id_to_name.get(style_id, style_id)
            role = STYLE_MAP.get(style_name, "description")

            if role == "heading":
                add_blank()
                md_lines.append(f"## {text}")
                add_blank()

            elif role == "subheading":
                add_blank()
                md_lines.append(f"### {text}")
                add_blank()

            elif role == "action":
                add_blank()
                md_lines.append(f"`{text}`")
                add_blank()

            elif role == "dialogue":
                # Keep consecutive dialogue lines together
                if md_lines and md_lines[-1] == "":
                    # Check if previous non-blank was also dialogue
                    prev = ""
                    for prev_line in reversed(md_lines):
                        if prev_line != "":
                            prev = prev_line
                            break
                    if prev.startswith(">") and not prev.startswith("> **"):
                        md_lines.append(">")
                md_lines.append(f"> {text}")

            elif role == "reflection":
                if md_lines and md_lines[-1] == "":
                    prev = ""
                    for prev_line in reversed(md_lines):
                        if prev_line != "":
                            prev = prev_line
                            break
                    if prev.startswith("> *"):
                        md_lines.append(">")
                md_lines.append(f"> *{text}*")

            else:  # description, list, default
                add_blank()
                md_lines.append(text)
                add_blank()

        elif tag == "tbl":
            add_blank()
            if is_reflection_table(element):
                lines = convert_reflection_table(element, style_id_to_name)
            else:
                lines = convert_data_table(element, style_id_to_name)
            md_lines.extend(lines)
            add_blank()

    # Clean up: collapse multiple blank lines
    cleaned = []
    for line in md_lines:
        if line == "" and cleaned and cleaned[-1] == "":
            continue
        cleaned.append(line)

    # Strip trailing blanks
    while cleaned and cleaned[-1] == "":
        cleaned.pop()

    output_text = "\n".join(cleaned) + "\n"
    output_path.write_text(output_text, encoding="utf-8")

    # Count stats
    n_images = len(rId_map)
    print(f"Converted: {docx_path}")
    print(f"Output:    {output_path}")
    print(f"Images:    {n_images} extracted to {media_dir}/")
    print(f"Lines:     {len(cleaned)}")

    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(f"Usage: python {sys.argv[0]} <input.docx> [output.md] [--media-dir DIR]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = None
    media_directory = None

    args = sys.argv[2:]
    i = 0
    while i < len(args):
        if args[i] == "--media-dir" and i + 1 < len(args):
            media_directory = args[i + 1]
            i += 2
        elif output_file is None:
            output_file = args[i]
            i += 1
        else:
            i += 1

    convert_docx_to_md(input_file, output_file, media_directory)
